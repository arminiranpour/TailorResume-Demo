from __future__ import annotations

import argparse
import io
import json
import logging
import unicodedata
from difflib import SequenceMatcher
from typing import Any, Dict, List, Tuple

from docx import Document
from docx.text.paragraph import Paragraph

from app.docx_engine.audit import AuditCollector, ReplacementAuditRecord
from app.docx_engine.types import (
    DocxAlignmentError,
    DocxInvariantError,
    DocxReplacementError,
    ReplacementTarget,
)

logger = logging.getLogger(__name__)

_BULLET_MARKERS = {
    "\u2022",
    "\u2023",
    "\u25E6",
    "\u2043",
    "\u2219",
    "\u00B7",
    "\u2027",
    "\u25AA",
    "\u25AB",
    "\u25CF",
    "\u25CB",
    "\u25A0",
    "\u25A1",
    "\u25C6",
    "\u25C7",
    "\u25B6",
    "\u25BA",
    "\u25B8",
    "\u25B9",
}


def get_resume_replacement_targets(
    original_resume_json: dict,
    tailored_resume_json: dict,
) -> List[Dict[str, Any]]:
    """Build an ordered list of replacement targets from original and tailored ResumeJSON."""
    original = _require_dict(original_resume_json, "original_resume_json")
    tailored = _require_dict(tailored_resume_json, "tailored_resume_json")

    _assert_resume_structure_matches(original, tailored)

    targets: List[ReplacementTarget] = []

    _append_summary_target(targets, original, tailored)
    _append_skill_targets(targets, original, tailored)
    _append_experience_targets(targets, original, tailored)
    _append_project_targets(targets, original, tailored)

    return [
        {
            "field_type": target.field_type,
            "field_id": target.field_id,
            "old_text": target.old_text,
            "new_text": target.new_text,
        }
        for target in targets
    ]


def get_paragraph_visible_text(paragraph: Paragraph) -> str:
    """Return concatenated text from all runs in a paragraph."""
    return "".join(run.text for run in paragraph.runs)


def normalize_for_verification(text: str) -> str:
    """Normalize text for deterministic paragraph alignment verification."""
    if text is None:
        return ""
    normalized = unicodedata.normalize("NFKC", str(text)).lower()
    for marker in _BULLET_MARKERS:
        normalized = normalized.replace(marker, " ")
    cleaned_chars: List[str] = []
    for char in normalized:
        if char.isalnum() or char.isspace():
            cleaned_chars.append(char)
        else:
            cleaned_chars.append(" ")
    normalized = "".join(cleaned_chars)
    normalized = " ".join(normalized.split())
    return normalized.strip()


def verify_paragraph_alignment(
    paragraph: Paragraph,
    expected_old_text: str,
    similarity_threshold: float = 0.92,
) -> Tuple[bool, float]:
    """Verify that a paragraph still matches the expected original text."""
    visible_text = get_paragraph_visible_text(paragraph)
    normalized_expected = normalize_for_verification(expected_old_text)
    normalized_visible = normalize_for_verification(visible_text)
    if normalized_expected == normalized_visible:
        return True, 1.0
    score = _similarity_score(normalized_expected, normalized_visible)
    return score >= similarity_threshold, score


def replace_paragraph_runs_preserving_structure(
    paragraph: Paragraph,
    new_text: str,
) -> Dict[str, Any]:
    """Replace paragraph text while preserving run structure."""
    old_run_count = len(paragraph.runs)
    warning: str | None = None

    if old_run_count == 0:
        warning = "Paragraph had no runs; created a single run for replacement."
        paragraph.add_run(new_text)
        return {
            "old_run_count": old_run_count,
            "new_run_count": len(paragraph.runs),
            "warning": warning,
        }

    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""

    if old_run_count > 1:
        warning = f"Paragraph had {old_run_count} runs; collapsed into first run."

    return {
        "old_run_count": old_run_count,
        "new_run_count": len(paragraph.runs),
        "warning": warning,
    }


def validate_replacement_invariants(
    original_doc: Document,
    edited_doc: Document,
    replaced_paragraph_indices: List[int],
) -> List[str]:
    """Validate that document invariants remain intact after replacement."""
    errors: List[str] = []
    original_count = len(original_doc.paragraphs)
    edited_count = len(edited_doc.paragraphs)

    if original_count != edited_count:
        errors.append(
            f"Paragraph count changed from {original_count} to {edited_count}."
        )

    for index in replaced_paragraph_indices:
        if index < 0 or index >= edited_count or index >= original_count:
            errors.append(f"Paragraph index {index} is out of bounds.")
            continue
        original_style = _style_signature(original_doc.paragraphs[index])
        edited_style = _style_signature(edited_doc.paragraphs[index])
        if original_style != edited_style:
            errors.append(
                f"Paragraph {index} style changed from {original_style} to {edited_style}."
            )
    return errors


def apply_tailored_text_to_docx(
    docx_path: str,
    original_resume_json: dict,
    tailored_resume_json: dict,
    docx_mapping: dict,
) -> Dict[str, Any]:
    """Apply tailored resume text to a DOCX using run-preserving replacement."""
    if not isinstance(docx_mapping, dict):
        raise DocxReplacementError(
            "docx_mapping must be a dict",
            details={"type": type(docx_mapping).__name__},
        )

    original_doc = Document(docx_path)
    edited_doc = Document(docx_path)

    targets = get_resume_replacement_targets(original_resume_json, tailored_resume_json)
    audit = AuditCollector(total_targets=len(targets))
    replaced_indices: List[int] = []

    for target in targets:
        field_type = target["field_type"]
        field_id = target["field_id"]
        old_text = target["old_text"]
        new_text = target["new_text"]

        paragraph_index = _resolve_mapping_index(docx_mapping, field_type, field_id)
        if paragraph_index is None:
            raise DocxReplacementError(
                f"Missing mapping for {field_type} {field_id}",
                details={"field_type": field_type, "field_id": field_id},
            )
        if paragraph_index < 0 or paragraph_index >= len(edited_doc.paragraphs):
            raise DocxReplacementError(
                f"Mapped paragraph index {paragraph_index} out of bounds",
                details={"field_type": field_type, "field_id": field_id},
            )

        paragraph = edited_doc.paragraphs[paragraph_index]
        ok, score = verify_paragraph_alignment(paragraph, old_text)
        if not ok:
            visible_text = get_paragraph_visible_text(paragraph)
            logger.error(
                "Alignment failed for %s %s at paragraph %s",
                field_type,
                field_id,
                paragraph_index,
            )
            raise DocxAlignmentError(
                f"Alignment failed for {field_type} {field_id} at paragraph {paragraph_index}",
                details={
                    "field_type": field_type,
                    "field_id": field_id,
                    "paragraph_index": paragraph_index,
                    "expected_text": old_text,
                    "visible_text": visible_text,
                    "verification_score": score,
                },
            )

        logger.info(
            "Replacing %s %s in paragraph %s",
            field_type,
            field_id,
            paragraph_index,
        )
        diagnostics = replace_paragraph_runs_preserving_structure(paragraph, new_text)
        warning = diagnostics.get("warning")
        if warning:
            logger.warning("Paragraph %s warning: %s", paragraph_index, warning)
            audit.add_warning(f"{field_type}:{field_id}: {warning}")

        audit.add_replacement(
            ReplacementAuditRecord(
                field_type=field_type,
                field_id=field_id,
                paragraph_index=paragraph_index,
                old_text=old_text,
                new_text=new_text,
                verification_score=score,
                old_run_count=diagnostics["old_run_count"],
                replacement_status="replaced",
                warning=warning,
            )
        )
        replaced_indices.append(paragraph_index)

    invariant_errors = validate_replacement_invariants(
        original_doc, edited_doc, replaced_indices
    )
    if invariant_errors:
        for error in invariant_errors:
            logger.error("Invariant violation: %s", error)
        raise DocxInvariantError(
            "Document invariants violated after replacement",
            details={"errors": invariant_errors},
        )

    buffer = io.BytesIO()
    edited_doc.save(buffer)
    return {"docx_bytes": buffer.getvalue(), "audit_log": audit.to_dict()}


def _similarity_score(a: str, b: str) -> float:
    if not a and not b:
        return 1.0
    if not a or not b:
        return 0.0
    tokens_a = set(a.split())
    tokens_b = set(b.split())
    union = tokens_a | tokens_b
    overlap = len(tokens_a & tokens_b) / len(union) if union else 0.0
    ratio = SequenceMatcher(None, a, b).ratio()
    return (overlap + ratio) / 2.0


def _style_signature(paragraph: Paragraph) -> Tuple[str | None, str | None]:
    style = paragraph.style
    if style is None:
        return (None, None)
    return (getattr(style, "style_id", None), getattr(style, "name", None))


def _resolve_mapping_index(
    docx_mapping: Dict[str, Any],
    field_type: str,
    field_id: str,
) -> int | None:
    if field_type == "summary":
        summary = docx_mapping.get("summary")
        return _extract_paragraph_index(summary)
    if field_type == "skills":
        skills = docx_mapping.get("skills", {})
        return _extract_paragraph_index(skills.get(field_id))
    if field_type in {"bullet", "project"}:
        bullets = docx_mapping.get("bullets", {})
        return _extract_paragraph_index(bullets.get(field_id))
    return None


def _extract_paragraph_index(value: Any) -> int | None:
    if value is None:
        return None
    if isinstance(value, int):
        return value
    if isinstance(value, dict) and isinstance(value.get("paragraph_index"), int):
        return value["paragraph_index"]
    return None


def _require_dict(value: Any, label: str) -> Dict[str, Any]:
    if not isinstance(value, dict):
        raise DocxReplacementError(
            f"{label} must be a dict",
            details={"type": type(value).__name__},
        )
    return value


def _assert_resume_structure_matches(
    original: Dict[str, Any],
    tailored: Dict[str, Any],
) -> None:
    errors: List[str] = []

    _compare_summary_structure(original, tailored, errors)
    _compare_skill_structure(original, tailored, errors)
    _compare_experience_structure(original, tailored, errors)
    _compare_project_structure(original, tailored, errors)

    if errors:
        raise DocxReplacementError(
            "Resume structure mismatch between original and tailored JSON",
            details={"errors": errors},
        )


def _compare_summary_structure(
    original: Dict[str, Any], tailored: Dict[str, Any], errors: List[str]
) -> None:
    original_summary = original.get("summary")
    tailored_summary = tailored.get("summary")
    if (original_summary is None) != (tailored_summary is None):
        errors.append("Summary presence differs between original and tailored JSON.")
        return
    if original_summary is None:
        return
    if not isinstance(original_summary, dict) or not isinstance(tailored_summary, dict):
        errors.append("Summary must be a dict in both original and tailored JSON.")
        return
    if original_summary.get("id") != tailored_summary.get("id"):
        errors.append("Summary id differs between original and tailored JSON.")


def _compare_skill_structure(
    original: Dict[str, Any], tailored: Dict[str, Any], errors: List[str]
) -> None:
    original_skills = original.get("skills")
    tailored_skills = tailored.get("skills")
    if (original_skills is None) != (tailored_skills is None):
        errors.append("Skills presence differs between original and tailored JSON.")
        return
    if original_skills is None:
        return
    if not isinstance(original_skills, dict) or not isinstance(tailored_skills, dict):
        errors.append("Skills must be dicts in both original and tailored JSON.")
        return
    original_lines = original_skills.get("lines")
    tailored_lines = tailored_skills.get("lines")
    if not isinstance(original_lines, list) or not isinstance(tailored_lines, list):
        errors.append("Skills lines must be lists in both original and tailored JSON.")
        return
    if len(original_lines) != len(tailored_lines):
        errors.append("Skills line count differs between original and tailored JSON.")
        return
    for index, (orig_line, tail_line) in enumerate(zip(original_lines, tailored_lines)):
        if not isinstance(orig_line, dict) or not isinstance(tail_line, dict):
            errors.append(f"Skills line {index} must be dicts in both JSONs.")
            continue
        if orig_line.get("line_id") != tail_line.get("line_id"):
            errors.append(f"Skills line id differs at index {index}.")


def _compare_experience_structure(
    original: Dict[str, Any], tailored: Dict[str, Any], errors: List[str]
) -> None:
    original_exp = original.get("experience")
    tailored_exp = tailored.get("experience")
    if not isinstance(original_exp, list) or not isinstance(tailored_exp, list):
        errors.append("Experience must be lists in both original and tailored JSON.")
        return
    if len(original_exp) != len(tailored_exp):
        errors.append("Experience count differs between original and tailored JSON.")
        return
    for exp_index, (orig_exp, tail_exp) in enumerate(zip(original_exp, tailored_exp)):
        if not isinstance(orig_exp, dict) or not isinstance(tail_exp, dict):
            errors.append(f"Experience entry {exp_index} must be dicts.")
            continue
        if orig_exp.get("exp_id") != tail_exp.get("exp_id"):
            errors.append(f"Experience id differs at index {exp_index}.")
        orig_bullets = orig_exp.get("bullets")
        tail_bullets = tail_exp.get("bullets")
        if not isinstance(orig_bullets, list) or not isinstance(tail_bullets, list):
            errors.append(f"Experience bullets must be lists at index {exp_index}.")
            continue
        if len(orig_bullets) != len(tail_bullets):
            errors.append(f"Bullet count differs for experience {exp_index}.")
            continue
        for bullet_index, (orig_bullet, tail_bullet) in enumerate(
            zip(orig_bullets, tail_bullets)
        ):
            if not isinstance(orig_bullet, dict) or not isinstance(tail_bullet, dict):
                errors.append(
                    f"Bullet {bullet_index} in experience {exp_index} must be dicts."
                )
                continue
            if orig_bullet.get("bullet_id") != tail_bullet.get("bullet_id"):
                errors.append(
                    f"Bullet id differs at experience {exp_index}, index {bullet_index}."
                )
            if orig_bullet.get("bullet_index") != tail_bullet.get("bullet_index"):
                errors.append(
                    "Bullet index value differs at "
                    f"experience {exp_index}, index {bullet_index}."
                )


def _compare_project_structure(
    original: Dict[str, Any], tailored: Dict[str, Any], errors: List[str]
) -> None:
    original_projects = original.get("projects")
    tailored_projects = tailored.get("projects")
    if original_projects is None and tailored_projects is None:
        return
    if (original_projects is None) != (tailored_projects is None):
        errors.append("Projects presence differs between original and tailored JSON.")
        return
    if not isinstance(original_projects, list) or not isinstance(tailored_projects, list):
        errors.append("Projects must be lists in both original and tailored JSON.")
        return
    if len(original_projects) != len(tailored_projects):
        errors.append("Project count differs between original and tailored JSON.")
        return
    for proj_index, (orig_proj, tail_proj) in enumerate(
        zip(original_projects, tailored_projects)
    ):
        if not isinstance(orig_proj, dict) or not isinstance(tail_proj, dict):
            errors.append(f"Project entry {proj_index} must be dicts.")
            continue
        if orig_proj.get("project_id") != tail_proj.get("project_id"):
            errors.append(f"Project id differs at index {proj_index}.")
        orig_bullets = orig_proj.get("bullets")
        tail_bullets = tail_proj.get("bullets")
        if not isinstance(orig_bullets, list) or not isinstance(tail_bullets, list):
            errors.append(f"Project bullets must be lists at index {proj_index}.")
            continue
        if len(orig_bullets) != len(tail_bullets):
            errors.append(f"Project bullet count differs at index {proj_index}.")
            continue
        for bullet_index, (orig_bullet, tail_bullet) in enumerate(
            zip(orig_bullets, tail_bullets)
        ):
            if not isinstance(orig_bullet, dict) or not isinstance(tail_bullet, dict):
                errors.append(
                    f"Project bullet {bullet_index} at index {proj_index} must be dicts."
                )
                continue
            if orig_bullet.get("bullet_id") != tail_bullet.get("bullet_id"):
                errors.append(
                    f"Project bullet id differs at index {proj_index}, bullet {bullet_index}."
                )
            if orig_bullet.get("bullet_index") != tail_bullet.get("bullet_index"):
                errors.append(
                    "Project bullet index value differs at "
                    f"index {proj_index}, bullet {bullet_index}."
                )


def _append_summary_target(
    targets: List[ReplacementTarget], original: Dict[str, Any], tailored: Dict[str, Any]
) -> None:
    summary = original.get("summary")
    tailored_summary = tailored.get("summary")
    if not isinstance(summary, dict) or not isinstance(tailored_summary, dict):
        return
    field_id = summary.get("id")
    if not isinstance(field_id, str):
        raise DocxReplacementError("Summary id must be a string")
    old_text = summary.get("text")
    new_text = tailored_summary.get("text")
    if not isinstance(old_text, str) or not isinstance(new_text, str):
        raise DocxReplacementError("Summary text must be strings")
    if old_text != new_text:
        targets.append(
            ReplacementTarget(
                field_type="summary",
                field_id=field_id,
                old_text=old_text,
                new_text=new_text,
            )
        )


def _append_skill_targets(
    targets: List[ReplacementTarget], original: Dict[str, Any], tailored: Dict[str, Any]
) -> None:
    original_skills = original.get("skills")
    tailored_skills = tailored.get("skills")
    if not isinstance(original_skills, dict) or not isinstance(tailored_skills, dict):
        return
    original_lines = original_skills.get("lines", [])
    tailored_lines = tailored_skills.get("lines", [])
    if not isinstance(original_lines, list) or not isinstance(tailored_lines, list):
        return
    for index, (orig_line, tail_line) in enumerate(zip(original_lines, tailored_lines)):
        if not isinstance(orig_line, dict) or not isinstance(tail_line, dict):
            raise DocxReplacementError(f"Skills line {index} must be dicts")
        line_id = orig_line.get("line_id")
        if not isinstance(line_id, str):
            raise DocxReplacementError("Skills line_id must be a string")
        old_text = orig_line.get("text")
        new_text = tail_line.get("text")
        if not isinstance(old_text, str) or not isinstance(new_text, str):
            raise DocxReplacementError("Skills line text must be strings")
        if old_text != new_text:
            targets.append(
                ReplacementTarget(
                    field_type="skills",
                    field_id=line_id,
                    old_text=old_text,
                    new_text=new_text,
                )
            )


def _append_experience_targets(
    targets: List[ReplacementTarget], original: Dict[str, Any], tailored: Dict[str, Any]
) -> None:
    original_exp = original.get("experience", [])
    tailored_exp = tailored.get("experience", [])
    if not isinstance(original_exp, list) or not isinstance(tailored_exp, list):
        return
    for exp_index, (orig_exp, tail_exp) in enumerate(zip(original_exp, tailored_exp)):
        if not isinstance(orig_exp, dict) or not isinstance(tail_exp, dict):
            raise DocxReplacementError(f"Experience entry {exp_index} must be dicts")
        orig_bullets = orig_exp.get("bullets", [])
        tail_bullets = tail_exp.get("bullets", [])
        if not isinstance(orig_bullets, list) or not isinstance(tail_bullets, list):
            raise DocxReplacementError(
                f"Experience bullets must be lists at index {exp_index}"
            )
        for bullet_index, (orig_bullet, tail_bullet) in enumerate(
            zip(orig_bullets, tail_bullets)
        ):
            if not isinstance(orig_bullet, dict) or not isinstance(
                tail_bullet, dict
            ):
                raise DocxReplacementError(
                    f"Bullet {bullet_index} in experience {exp_index} must be dicts"
                )
            bullet_id = orig_bullet.get("bullet_id")
            if not isinstance(bullet_id, str):
                raise DocxReplacementError("Bullet id must be a string")
            old_text = orig_bullet.get("text")
            new_text = tail_bullet.get("text")
            if not isinstance(old_text, str) or not isinstance(new_text, str):
                raise DocxReplacementError("Bullet text must be strings")
            if old_text != new_text:
                targets.append(
                    ReplacementTarget(
                        field_type="bullet",
                        field_id=bullet_id,
                        old_text=old_text,
                        new_text=new_text,
                    )
                )


def _append_project_targets(
    targets: List[ReplacementTarget], original: Dict[str, Any], tailored: Dict[str, Any]
) -> None:
    original_projects = original.get("projects", [])
    tailored_projects = tailored.get("projects", [])
    if not isinstance(original_projects, list) or not isinstance(
        tailored_projects, list
    ):
        return
    for proj_index, (orig_proj, tail_proj) in enumerate(
        zip(original_projects, tailored_projects)
    ):
        if not isinstance(orig_proj, dict) or not isinstance(tail_proj, dict):
            raise DocxReplacementError(f"Project entry {proj_index} must be dicts")
        orig_bullets = orig_proj.get("bullets", [])
        tail_bullets = tail_proj.get("bullets", [])
        if not isinstance(orig_bullets, list) or not isinstance(tail_bullets, list):
            raise DocxReplacementError(
                f"Project bullets must be lists at index {proj_index}"
            )
        for bullet_index, (orig_bullet, tail_bullet) in enumerate(
            zip(orig_bullets, tail_bullets)
        ):
            if not isinstance(orig_bullet, dict) or not isinstance(
                tail_bullet, dict
            ):
                raise DocxReplacementError(
                    f"Project bullet {bullet_index} at index {proj_index} must be dicts"
                )
            bullet_id = orig_bullet.get("bullet_id")
            if not isinstance(bullet_id, str):
                raise DocxReplacementError("Project bullet id must be a string")
            old_text = orig_bullet.get("text")
            new_text = tail_bullet.get("text")
            if not isinstance(old_text, str) or not isinstance(new_text, str):
                raise DocxReplacementError("Project bullet text must be strings")
            if old_text != new_text:
                targets.append(
                    ReplacementTarget(
                        field_type="project",
                        field_id=bullet_id,
                        old_text=old_text,
                        new_text=new_text,
                    )
                )


def _load_json(path: str) -> Dict[str, Any]:
    with open(path, "r", encoding="utf-8") as handle:
        data = json.load(handle)
    if not isinstance(data, dict):
        raise ValueError(f"Expected JSON object in {path}")
    return data


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Apply tailored ResumeJSON text to a DOCX using run-preserving replacement."
    )
    parser.add_argument("docx_path", help="Path to the original DOCX file")
    parser.add_argument("original_json", help="Path to original ResumeJSON")
    parser.add_argument("tailored_json", help="Path to tailored ResumeJSON")
    parser.add_argument("mapping_json", help="Path to DOCX mapping JSON")
    parser.add_argument(
        "--output",
        default="edited.docx",
        help="Output path for the edited DOCX",
    )
    args = parser.parse_args()

    original_json = _load_json(args.original_json)
    tailored_json = _load_json(args.tailored_json)
    mapping_json = _load_json(args.mapping_json)

    result = apply_tailored_text_to_docx(
        args.docx_path, original_json, tailored_json, mapping_json
    )

    with open(args.output, "wb") as handle:
        handle.write(result["docx_bytes"])

    print(json.dumps(result["audit_log"], indent=2))
