from __future__ import annotations

import logging
from typing import Any, Dict, List

from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph

from app.docx_engine.mapping import normalize_text
from app.docx_engine.types import DocxOverflowError

logger = logging.getLogger(__name__)


def get_visible_paragraph_text(paragraph: Paragraph) -> str:
    """Return concatenated visible text from all runs in a paragraph."""
    return "".join(run.text for run in paragraph.runs)


def collect_editable_field_records(
    doc: DocxDocument,
    resume_json: Dict[str, Any],
    docx_mapping: Dict[str, Any],
) -> List[Dict[str, Any]]:
    """Collect per-field metrics for all editable ResumeJSON fields."""
    if not isinstance(resume_json, dict):
        raise DocxOverflowError(
            "resume_json must be a dict",
            details={"type": type(resume_json).__name__},
        )
    if not isinstance(docx_mapping, dict):
        raise DocxOverflowError(
            "docx_mapping must be a dict",
            details={"type": type(docx_mapping).__name__},
        )

    fields = _extract_resume_fields(resume_json)
    records: List[Dict[str, Any]] = []
    paragraph_count = len(doc.paragraphs)

    for field in fields:
        paragraph_index = _resolve_mapping_index(
            docx_mapping, field["field_type"], field["field_id"]
        )
        if paragraph_index is None:
            raise DocxOverflowError(
                f"Missing mapping for {field['field_type']} {field['field_id']}",
                details={
                    "field_type": field["field_type"],
                    "field_id": field["field_id"],
                },
            )
        if paragraph_index < 0 or paragraph_index >= paragraph_count:
            raise DocxOverflowError(
                f"Mapped paragraph index {paragraph_index} out of bounds",
                details={
                    "field_type": field["field_type"],
                    "field_id": field["field_id"],
                    "paragraph_index": paragraph_index,
                    "paragraph_count": paragraph_count,
                },
            )

        paragraph = doc.paragraphs[paragraph_index]
        visible_text = get_visible_paragraph_text(paragraph)
        normalized_text = normalize_text(visible_text)
        style_name = getattr(paragraph.style, "name", None)

        record = {
            "field_type": field["field_type"],
            "field_id": field["field_id"],
            "paragraph_index": paragraph_index,
            "expected_text": field["text"],
            "visible_text": visible_text,
            "visible_length": len(visible_text),
            "normalized_length": len(normalized_text),
            "run_count": len(paragraph.runs),
            "paragraph_style_name": style_name,
        }
        records.append(record)

    logger.info("Collected %s editable field records", len(records))
    return records


def compute_docx_metrics(
    doc: DocxDocument,
    resume_json: Dict[str, Any],
    docx_mapping: Dict[str, Any],
) -> Dict[str, Any]:
    """Compute deterministic metrics for a DOCX document."""
    field_records = collect_editable_field_records(doc, resume_json, docx_mapping)

    total_editable_chars = sum(record["visible_length"] for record in field_records)
    total_editable_normalized_chars = sum(
        record["normalized_length"] for record in field_records
    )
    max_paragraph_visible_length = (
        max((record["visible_length"] for record in field_records), default=0)
    )
    sum_paragraph_visible_length = total_editable_chars

    metrics = {
        "field_records": field_records,
        "totals": {
            "total_editable_chars": total_editable_chars,
            "total_editable_normalized_chars": total_editable_normalized_chars,
            "max_paragraph_visible_length": max_paragraph_visible_length,
            "sum_paragraph_visible_length": sum_paragraph_visible_length,
            "editable_field_count": len(field_records),
        },
        "paragraph_count": len(doc.paragraphs),
    }

    logger.info(
        "Computed total editable chars: %s (normalized %s)",
        total_editable_chars,
        total_editable_normalized_chars,
    )
    return metrics


def _extract_resume_fields(resume_json: Dict[str, Any]) -> List[Dict[str, Any]]:
    fields: List[Dict[str, Any]] = []

    summary = resume_json.get("summary")
    if isinstance(summary, dict):
        summary_text = summary.get("text")
        if isinstance(summary_text, str) and summary_text.strip():
            summary_id = summary.get("id")
            fields.append(
                {
                    "field_type": "summary",
                    "field_id": summary_id if isinstance(summary_id, str) else "summary",
                    "text": summary_text,
                }
            )

    skills = resume_json.get("skills")
    if isinstance(skills, dict):
        lines = skills.get("lines")
    elif isinstance(skills, list):
        lines = skills
    else:
        lines = []
    if isinstance(lines, list):
        for idx, line in enumerate(lines):
            if not isinstance(line, dict):
                continue
            line_text = line.get("text")
            if not isinstance(line_text, str) or not line_text.strip():
                continue
            line_id = line.get("line_id")
            fields.append(
                {
                    "field_type": "skills",
                    "field_id": line_id if isinstance(line_id, str) else f"skills_{idx}",
                    "text": line_text,
                }
            )

    _extend_bullet_fields(fields, resume_json.get("experience"), "experience", "bullet")
    _extend_bullet_fields(fields, resume_json.get("projects"), "projects", "project")

    return fields


def _extend_bullet_fields(
    fields: List[Dict[str, Any]],
    collection: Any,
    label: str,
    field_type: str,
) -> None:
    if not isinstance(collection, list):
        return
    for exp_idx, item in enumerate(collection):
        if not isinstance(item, dict):
            continue
        bullets = item.get("bullets") if isinstance(item.get("bullets"), list) else []
        for bullet_idx, bullet in enumerate(bullets):
            if not isinstance(bullet, dict):
                continue
            bullet_text = bullet.get("text")
            if not isinstance(bullet_text, str) or not bullet_text.strip():
                continue
            bullet_id = bullet.get("bullet_id")
            fields.append(
                {
                    "field_type": field_type,
                    "field_id": (
                        bullet_id
                        if isinstance(bullet_id, str)
                        else f"{label}_{exp_idx}_b{bullet_idx}"
                    ),
                    "text": bullet_text,
                }
            )


def _resolve_mapping_index(
    docx_mapping: Dict[str, Any], field_type: str, field_id: str
) -> int | None:
    if field_type == "summary":
        summary = docx_mapping.get("summary")
        return _extract_paragraph_index(summary)
    if field_type == "skills":
        skills = docx_mapping.get("skills", {})
        return _extract_paragraph_index(skills.get(field_id))
    if field_type in {"bullet", "project"}:
        bullets = docx_mapping.get("bullets", {})
        return _extract_paragraph_index(bullets.get(field_id))
    return None


def _extract_paragraph_index(value: Any) -> int | None:
    if value is None:
        return None
    if isinstance(value, int):
        return value
    if isinstance(value, dict) and isinstance(value.get("paragraph_index"), int):
        return value["paragraph_index"]
    return None
