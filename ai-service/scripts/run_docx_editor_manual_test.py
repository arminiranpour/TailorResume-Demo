import io
import json
import os
import tempfile

from docx import Document

from app.docx_engine.editor import apply_tailored_text_to_docx
from app.docx_engine.mapping import build_docx_mapping


def _load_fixture_json(filename: str) -> dict:
    root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    path = os.path.join(root, "tests", "fixtures", filename)
    with open(path, "r", encoding="utf-8") as handle:
        return json.load(handle)


def _add_paragraph_with_runs(doc: Document, text: str, *, style: str | None = None) -> None:
    if style:
        paragraph = doc.add_paragraph(style=style)
    else:
        paragraph = doc.add_paragraph()
    if " " in text:
        first, rest = text.split(" ", 1)
        paragraph.add_run(f"{first} ")
        paragraph.add_run(rest)
    else:
        paragraph.add_run(text)


def create_temp_docx(original_resume_json: dict, output_path: str) -> None:
    doc = Document()
    doc.add_paragraph(original_resume_json["summary"]["text"])

    for line in original_resume_json["skills"]["lines"]:
        doc.add_paragraph(line["text"])

    bullets = original_resume_json["experience"][0]["bullets"]
    _add_paragraph_with_runs(doc, bullets[0]["text"], style="List Bullet")
    doc.add_paragraph(bullets[1]["text"], style="List Bullet")

    doc.save(output_path)


def main() -> None:
    original_json = _load_fixture_json("original_resume.json")
    tailored_json = _load_fixture_json("tailored_resume_one_change.json")

    with tempfile.TemporaryDirectory() as tmp_dir:
        original_docx = os.path.join(tmp_dir, "original_resume.docx")
        create_temp_docx(original_json, original_docx)

        mapping = build_docx_mapping(original_docx, original_json)
        result = apply_tailored_text_to_docx(
            original_docx, original_json, tailored_json, mapping
        )

        edited_path = os.path.join(tmp_dir, "edited_resume.docx")
        with open(edited_path, "wb") as handle:
            handle.write(result["docx_bytes"])

        print("Edited DOCX saved to:", edited_path)
        print("Audit summary:", json.dumps(result["audit_log"]["summary"], indent=2))
        print("Mapped paragraph indices:", json.dumps(mapping, indent=2))


if __name__ == "__main__":
    main()
