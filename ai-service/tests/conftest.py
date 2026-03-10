import json
import os
import sys
from copy import deepcopy
from typing import Dict

import pytest
from docx import Document


ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
if ROOT not in sys.path:
    sys.path.insert(0, ROOT)

FIXTURES_DIR = os.path.join(os.path.dirname(__file__), "fixtures")


def _load_fixture_json(filename: str) -> Dict:
    path = os.path.join(FIXTURES_DIR, filename)
    with open(path, "r", encoding="utf-8") as handle:
        return json.load(handle)


def _add_paragraph(doc: Document, text: str, *, style: str | None = None) -> None:
    if style:
        doc.add_paragraph(text, style=style)
    else:
        doc.add_paragraph(text)


def _add_paragraph_with_runs(
    doc: Document, text: str, *, style: str | None = None
) -> None:
    if style:
        paragraph = doc.add_paragraph(style=style)
    else:
        paragraph = doc.add_paragraph()
    if " " in text:
        first, rest = text.split(" ", 1)
        paragraph.add_run(f"{first} ")
        paragraph.add_run(rest)
    else:
        paragraph.add_run(text)


@pytest.fixture
def original_resume_json() -> Dict:
    return _load_fixture_json("original_resume.json")


@pytest.fixture
def tailored_resume_one_change_json() -> Dict:
    return _load_fixture_json("tailored_resume_one_change.json")


@pytest.fixture
def tailored_resume_bad_structure_json() -> Dict:
    return _load_fixture_json("tailored_resume_bad_structure.json")


@pytest.fixture
def tailored_resume_large_growth_json(original_resume_json) -> Dict:
    data = deepcopy(original_resume_json)
    bullet = data["experience"][0]["bullets"][0]
    extra = " Delivered scalable payment infrastructure improvements."
    bullet["text"] = f"{bullet['text']}{extra * 4}"
    bullet["char_count"] = len(bullet["text"])
    return data


@pytest.fixture
def temp_docx_path(tmp_path, original_resume_json) -> str:
    doc = Document()

    summary_text = original_resume_json["summary"]["text"]
    _add_paragraph(doc, summary_text)

    skill_lines = original_resume_json["skills"]["lines"]
    for line in skill_lines:
        _add_paragraph(doc, line["text"])

    bullets = original_resume_json["experience"][0]["bullets"]
    _add_paragraph_with_runs(doc, bullets[0]["text"], style="List Bullet")
    _add_paragraph(doc, bullets[1]["text"], style="List Bullet")

    path = tmp_path / "original_resume.docx"
    doc.save(path)
    return str(path)
