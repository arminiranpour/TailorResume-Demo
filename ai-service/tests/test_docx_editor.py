import io
from copy import deepcopy

import pytest
from docx import Document

from app.docx_engine.editor import (
    apply_tailored_text_to_docx,
    get_paragraph_visible_text,
    get_resume_replacement_targets,
    validate_replacement_invariants,
    verify_paragraph_alignment,
)
from app.docx_engine.mapping import build_docx_mapping
from app.docx_engine.types import DocxAlignmentError, DocxReplacementError


def _resolve_index(value):
    if isinstance(value, int):
        return value
    if isinstance(value, dict) and isinstance(value.get("paragraph_index"), int):
        return value["paragraph_index"]
    return None


def test_get_resume_replacement_targets_only_returns_changed_fields(
    original_resume_json, tailored_resume_one_change_json
):
    targets = get_resume_replacement_targets(
        original_resume_json, tailored_resume_one_change_json
    )
    assert len(targets) == 1
    target = targets[0]
    assert target["field_type"] == "bullet"
    assert target["field_id"] == "exp_1_b1"
    assert target["old_text"] == "Built APIs for payments."
    assert target["new_text"] == "Built secure payment APIs."


def test_apply_tailored_text_to_docx_happy_path(
    temp_docx_path, original_resume_json, tailored_resume_one_change_json
):
    mapping = build_docx_mapping(temp_docx_path, original_resume_json)
    result = apply_tailored_text_to_docx(
        temp_docx_path,
        original_resume_json,
        tailored_resume_one_change_json,
        mapping,
    )

    assert "docx_bytes" in result
    assert "audit_log" in result

    edited_doc = Document(io.BytesIO(result["docx_bytes"]))
    original_doc = Document(temp_docx_path)
    assert len(edited_doc.paragraphs) == len(original_doc.paragraphs)

    bullet_index = _resolve_index(mapping["bullets"]["exp_1_b1"])
    assert bullet_index is not None

    assert (
        get_paragraph_visible_text(edited_doc.paragraphs[bullet_index])
        == "Built secure payment APIs."
    )

    for idx, paragraph in enumerate(original_doc.paragraphs):
        if idx == bullet_index:
            continue
        assert get_paragraph_visible_text(paragraph) == get_paragraph_visible_text(
            edited_doc.paragraphs[idx]
        )


def test_apply_tailored_text_to_docx_rejects_structure_mismatch(
    temp_docx_path, original_resume_json, tailored_resume_bad_structure_json
):
    mapping = build_docx_mapping(temp_docx_path, original_resume_json)
    with pytest.raises(DocxReplacementError):
        apply_tailored_text_to_docx(
            temp_docx_path,
            original_resume_json,
            tailored_resume_bad_structure_json,
            mapping,
        )


def test_apply_tailored_text_to_docx_rejects_bad_mapping(
    temp_docx_path, original_resume_json, tailored_resume_one_change_json
):
    mapping = build_docx_mapping(temp_docx_path, original_resume_json)
    bad_mapping = deepcopy(mapping)
    bad_mapping["bullets"]["exp_1_b1"] = 0
    with pytest.raises(DocxAlignmentError):
        apply_tailored_text_to_docx(
            temp_docx_path,
            original_resume_json,
            tailored_resume_one_change_json,
            bad_mapping,
        )


def test_verify_paragraph_alignment_fails_on_wrong_paragraph(temp_docx_path):
    doc = Document(temp_docx_path)
    ok, score = verify_paragraph_alignment(doc.paragraphs[0], "Completely different")
    assert ok is False
    assert score < 0.92


def test_validate_replacement_invariants_preserves_paragraph_count_and_styles(
    temp_docx_path, original_resume_json, tailored_resume_one_change_json
):
    mapping = build_docx_mapping(temp_docx_path, original_resume_json)
    result = apply_tailored_text_to_docx(
        temp_docx_path,
        original_resume_json,
        tailored_resume_one_change_json,
        mapping,
    )
    original_doc = Document(temp_docx_path)
    edited_doc = Document(io.BytesIO(result["docx_bytes"]))

    bullet_index = _resolve_index(mapping["bullets"]["exp_1_b1"])
    errors = validate_replacement_invariants(
        original_doc, edited_doc, [bullet_index]
    )
    assert errors == []

    assert (
        original_doc.paragraphs[bullet_index].style.name
        == edited_doc.paragraphs[bullet_index].style.name
    )


def test_run_preserving_replacement_keeps_first_run_and_clears_rest(
    temp_docx_path, original_resume_json, tailored_resume_one_change_json
):
    mapping = build_docx_mapping(temp_docx_path, original_resume_json)
    result = apply_tailored_text_to_docx(
        temp_docx_path,
        original_resume_json,
        tailored_resume_one_change_json,
        mapping,
    )
    edited_doc = Document(io.BytesIO(result["docx_bytes"]))
    original_doc = Document(temp_docx_path)

    bullet_index = _resolve_index(mapping["bullets"]["exp_1_b1"])
    original_run_count = len(original_doc.paragraphs[bullet_index].runs)
    assert original_run_count >= 2

    runs = edited_doc.paragraphs[bullet_index].runs
    assert runs[0].text == "Built secure payment APIs."
    for run in runs[1:]:
        assert run.text == ""


def test_audit_log_is_deterministic_and_complete(
    temp_docx_path, original_resume_json, tailored_resume_one_change_json
):
    mapping = build_docx_mapping(temp_docx_path, original_resume_json)

    result_a = apply_tailored_text_to_docx(
        temp_docx_path,
        original_resume_json,
        tailored_resume_one_change_json,
        mapping,
    )
    result_b = apply_tailored_text_to_docx(
        temp_docx_path,
        original_resume_json,
        tailored_resume_one_change_json,
        mapping,
    )

    assert result_a["audit_log"] == result_b["audit_log"]

    audit_log = result_a["audit_log"]
    assert audit_log["summary"]["total_targets"] == 1
    assert audit_log["summary"]["replaced"] == 1
    assert audit_log["summary"]["failed"] == 0

    expected_keys = {
        "field_type",
        "field_id",
        "paragraph_index",
        "old_text",
        "new_text",
        "verification_score",
        "old_run_count",
        "replacement_status",
        "warning",
    }
    for record in audit_log["replacements"]:
        assert expected_keys.issubset(record.keys())


def test_output_docx_bytes_can_be_reopened(
    temp_docx_path, original_resume_json, tailored_resume_one_change_json, tmp_path
):
    mapping = build_docx_mapping(temp_docx_path, original_resume_json)
    result = apply_tailored_text_to_docx(
        temp_docx_path,
        original_resume_json,
        tailored_resume_one_change_json,
        mapping,
    )
    output_path = tmp_path / "edited.docx"
    output_path.write_bytes(result["docx_bytes"])

    reopened = Document(str(output_path))
    assert len(reopened.paragraphs) > 0
