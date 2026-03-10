"""Overflow heuristic tests.

Run locally:
source .venv/bin/activate
PYTHONPATH=. python -m pytest tests/test_docx_editor.py tests/test_docx_overflow.py -q
"""

import io

import pytest
from docx import Document

from app.docx_engine.editor import apply_tailored_text_to_docx
from app.docx_engine.mapping import build_docx_mapping
from app.docx_engine.overflow import evaluate_docx_overflow_risk
from app.docx_engine.types import DocxOverflowError


def test_overflow_check_passes_for_small_safe_change(
    temp_docx_path, original_resume_json, tailored_resume_one_change_json
):
    mapping = build_docx_mapping(temp_docx_path, original_resume_json)
    result = apply_tailored_text_to_docx(
        temp_docx_path,
        original_resume_json,
        tailored_resume_one_change_json,
        mapping,
    )

    report = evaluate_docx_overflow_risk(
        temp_docx_path,
        result["docx_bytes"],
        original_resume_json,
        tailored_resume_one_change_json,
        mapping,
    )

    assert report["valid"] is True
    assert report["overflow_risk"] is False
    assert report["rules_failed"] == []


def test_overflow_check_flags_large_growth(
    temp_docx_path, original_resume_json, tailored_resume_large_growth_json
):
    mapping = build_docx_mapping(temp_docx_path, original_resume_json)
    result = apply_tailored_text_to_docx(
        temp_docx_path,
        original_resume_json,
        tailored_resume_large_growth_json,
        mapping,
    )

    report = evaluate_docx_overflow_risk(
        temp_docx_path,
        result["docx_bytes"],
        original_resume_json,
        tailored_resume_large_growth_json,
        mapping,
    )

    assert report["overflow_risk"] is True
    assert report["rules_failed"]

    largest = report["largest_deltas"][0]
    assert largest["field_id"] == "exp_1_b1"


def test_overflow_threshold_behavior_is_deterministic(
    temp_docx_path, original_resume_json, tailored_resume_one_change_json
):
    mapping = build_docx_mapping(temp_docx_path, original_resume_json)
    result = apply_tailored_text_to_docx(
        temp_docx_path,
        original_resume_json,
        tailored_resume_one_change_json,
        mapping,
    )

    thresholds = {
        "max_total_growth_chars": 2,
        "max_total_growth_percent": 0.03,
        "max_field_growth_chars": 2,
        "max_field_growth_percent": 0.10,
        "max_longest_paragraph_growth_chars": 2,
        "max_longest_paragraph_growth_percent": 0.10,
        "max_increased_paragraphs": 1,
    }

    report_a = evaluate_docx_overflow_risk(
        temp_docx_path,
        result["docx_bytes"],
        original_resume_json,
        tailored_resume_one_change_json,
        mapping,
        thresholds=thresholds,
    )
    report_b = evaluate_docx_overflow_risk(
        temp_docx_path,
        result["docx_bytes"],
        original_resume_json,
        tailored_resume_one_change_json,
        mapping,
        thresholds=thresholds,
    )

    assert report_a == report_b
    assert report_a["rules_failed"] == report_b["rules_failed"]


def test_overflow_report_shape_is_complete(
    temp_docx_path, original_resume_json, tailored_resume_one_change_json
):
    mapping = build_docx_mapping(temp_docx_path, original_resume_json)
    result = apply_tailored_text_to_docx(
        temp_docx_path,
        original_resume_json,
        tailored_resume_one_change_json,
        mapping,
    )
    report = evaluate_docx_overflow_risk(
        temp_docx_path,
        result["docx_bytes"],
        original_resume_json,
        tailored_resume_one_change_json,
        mapping,
    )

    for key in ["summary", "field_reports", "largest_deltas", "rules_failed", "valid", "overflow_risk"]:
        assert key in report


def test_editor_and_overflow_preserve_paragraph_count(
    temp_docx_path, original_resume_json, tailored_resume_one_change_json
):
    mapping = build_docx_mapping(temp_docx_path, original_resume_json)
    result = apply_tailored_text_to_docx(
        temp_docx_path,
        original_resume_json,
        tailored_resume_one_change_json,
        mapping,
    )
    report = evaluate_docx_overflow_risk(
        temp_docx_path,
        result["docx_bytes"],
        original_resume_json,
        tailored_resume_one_change_json,
        mapping,
    )

    original_doc = Document(temp_docx_path)
    edited_doc = Document(io.BytesIO(result["docx_bytes"]))

    assert report["summary"]["paragraph_count_original"] == len(original_doc.paragraphs)
    assert report["summary"]["paragraph_count_edited"] == len(edited_doc.paragraphs)
    assert report["summary"]["paragraph_count_original"] == report["summary"]["paragraph_count_edited"]


def test_editor_and_overflow_preserve_replaced_paragraph_styles(
    temp_docx_path, original_resume_json, tailored_resume_one_change_json
):
    mapping = build_docx_mapping(temp_docx_path, original_resume_json)
    result = apply_tailored_text_to_docx(
        temp_docx_path,
        original_resume_json,
        tailored_resume_one_change_json,
        mapping,
    )
    report = evaluate_docx_overflow_risk(
        temp_docx_path,
        result["docx_bytes"],
        original_resume_json,
        tailored_resume_one_change_json,
        mapping,
    )

    original_doc = Document(temp_docx_path)
    edited_doc = Document(io.BytesIO(result["docx_bytes"]))

    bullet_index = mapping["bullets"]["exp_1_b1"]
    if isinstance(bullet_index, dict):
        bullet_index = bullet_index["paragraph_index"]

    assert (
        original_doc.paragraphs[bullet_index].style.name
        == edited_doc.paragraphs[bullet_index].style.name
    )

    assert report["summary"]["replaced_paragraph_style_names"] == [
        edited_doc.paragraphs[bullet_index].style.name
    ]


def test_overflow_evaluator_rejects_invalid_inputs(
    temp_docx_path, original_resume_json, tailored_resume_bad_structure_json
):
    mapping = build_docx_mapping(temp_docx_path, original_resume_json)
    result = apply_tailored_text_to_docx(
        temp_docx_path,
        original_resume_json,
        original_resume_json,
        mapping,
    )

    with pytest.raises(DocxOverflowError):
        evaluate_docx_overflow_risk(
            temp_docx_path,
            result["docx_bytes"],
            original_resume_json,
            tailored_resume_bad_structure_json,
            mapping,
        )
